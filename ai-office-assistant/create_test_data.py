"""
一键生成项目测试所需的所有文档（Excel + Word）
运行后会自动在 data/ 文件夹下创建销售表、模板、通讯录等文件
如果文件已存在，将被覆盖
"""

import os
import pandas as pd
from docx import Document

# 确保 data 文件夹存在
os.makedirs("data", exist_ok=True)

# ========== 1. 销售表A ==========
df_a = pd.DataFrame({
    "区域": ["华北", "华东"],
    "销售额": [15000, 12000],
    "负责人": ["张三", "李四"]
})
df_a.to_excel("data/销售表A.xlsx", index=False, engine="openpyxl")
print("✅ 已生成 data/销售表A.xlsx")

# ========== 2. 销售表B ==========
df_b = pd.DataFrame({
    "区域": ["华南", "华北"],
    "销售额": [18000, 20000],
    "负责人": ["王五", "张三"]
})
df_b.to_excel("data/销售表B.xlsx", index=False, engine="openpyxl")
print("✅ 已生成 data/销售表B.xlsx")

# ========== 3. 销售表C ==========
df_c = pd.DataFrame({
    "区域": ["华东", "西南"],
    "销售额": [9000, 11000],
    "负责人": ["李四", "赵六"]
})
df_c.to_excel("data/销售表C.xlsx", index=False, engine="openpyxl")
print("✅ 已生成 data/销售表C.xlsx")

# ========== 4. 填充数据 ==========
df_fill = pd.DataFrame({
    "姓名": ["张三", "李四", "王五"],
    "部门": ["华北", "华东", "华南"],
    "本周销售额": [35000, 21000, 18000],
    "完成率": ["95%", "88%", "92%"]
})
df_fill.to_excel("data/填充数据.xlsx", index=False, engine="openpyxl")
print("✅ 已生成 data/填充数据.xlsx")

# ========== 5. 通讯录 ==========
df_contacts = pd.DataFrame({
    "姓名": ["张三", "李四"],
    "邮箱": ["2582737926@qq.com", "2582737926@qq.com"]  # 请替换成你自己的邮箱
})
df_contacts.to_excel("data/通讯录.xlsx", index=False, engine="openpyxl")
print("✅ 已生成 data/通讯录.xlsx（记得修改邮箱地址！）")

# ========== 6. Word模板 ==========
doc = Document()
doc.add_paragraph("销售周报")
doc.add_paragraph("")
doc.add_paragraph("姓名：{{姓名}}")
doc.add_paragraph("部门：{{部门}}")
doc.add_paragraph("本周销售额：{{本周销售额}}")
doc.add_paragraph("完成率：{{完成率}}")
doc.add_paragraph("")
doc.add_paragraph("备注：以上数据由系统自动生成，请核对。")
doc.save("data/模板.docx")
print("✅ 已生成 data/模板.docx")

print("\n🎉 所有文件已生成完毕，放在 data/ 文件夹下。")
print("⚠️  请打开 data/通讯录.xlsx 把邮箱替换成你自己的真实邮箱后再测试邮件功能。")