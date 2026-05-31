"""
Word自动化处理模块
基于模板占位符替换，实现文档批量生成
"""

import os
import pandas as pd
from docx import Document


def fill_word_template(template_path, data_excel_path, output_folder="./output"):
    """
    根据Excel数据批量填充Word模板
    参数：template_path - Word模板文件路径
          data_excel_path - 包含填充数据的Excel文件路径
          output_folder - 输出文件夹路径
    返回：list - 生成的文件路径列表
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"模板文件不存在：{template_path}")
    if not os.path.exists(data_excel_path):
        raise FileNotFoundError(f"数据文件不存在：{data_excel_path}")

    os.makedirs(output_folder, exist_ok=True)
    data_df = pd.read_excel(data_excel_path, engine='openpyxl')  # 修复点
    generated_files = []

    for _, row in data_df.iterrows():
        doc = Document(template_path)

        # 替换段落中的占位符
        for paragraph in doc.paragraphs:
            for col in data_df.columns:
                placeholder = f"{{{{{col}}}}}"
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(row[col]))

        # 替换表格中的占位符
        for table in doc.tables:
            for row_cells in table.rows:
                for cell in row_cells.cells:
                    for paragraph in cell.paragraphs:
                        for col in data_df.columns:
                            placeholder = f"{{{{{col}}}}}"
                            if placeholder in paragraph.text:
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str(row[col]))

        filename = str(row.iloc[0])
        output_path = os.path.join(output_folder, f"{filename}_文档.docx")
        doc.save(output_path)
        generated_files.append(output_path)

    return generated_files